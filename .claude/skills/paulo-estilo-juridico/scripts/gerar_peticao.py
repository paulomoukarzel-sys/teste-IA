#!/usr/bin/env python3
"""
Gera peticoes juridicas formatadas em .docx no padrao do escritorio
Paulo Ekke Moukarzel Junior.

Template oficial: C:\\Users\\paulo\\OneDrive\\Documentos\\M.Adv\\Modelos\\Peticao inicial - com timbre Paulo.dotx
(timbre no cabecalho, dados do escritorio no rodape)

Salva SEMPRE em: <base-dir>/<CLIENTE>/output_claude/<TIPO>_<CLIENTE>_<DATA>.docx
"""

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print(
        "Erro: python-docx nao instalado.\n"
        "Instale com: pip install python-docx",
        file=__import__('sys').stderr,
    )
    __import__('sys').exit(1)

import argparse
import os
import re
import shutil
import sys
import tempfile
import unicodedata
import zipfile
from datetime import datetime

# Caminho do template oficial (com timbre Paulo)
_TEMPLATE_PATH = r'C:\Users\paulo\OneDrive\Documentos\M.Adv\Modelos\Petição inicial - com timbre Paulo.dotx'

# ── Regex de classificacao de paragrafos ─────────────────────────────────────
_RE_FECHAMENTO = re.compile(
    r'^(Pede deferimento|Florianópolis|Paulo Ekke|OAB/)',
    re.IGNORECASE,
)
_RE_NAO_NUMERADO = re.compile(
    r'^(Pleiteia,\s*por\s*fim|Por\s*fim,\s*sejam|Ante\s*o\s*exposto|'
    r'Diante\s*do\s*exposto|Requer,?\s*desde\s*já|Requer,|'
    r'Autos\s*n\.|REsp\s*n\.|RE\s*n\.|AREsp\s*n\.|'
    r'Recorrente:|Recorrido:|DOCUMENTOS\s*QUE|_{5,})',
    re.IGNORECASE,
)
# Linhas que iniciam a Parte 2 (Razões Recursais) → recebem quebra de página antes
_RE_INICIO_PARTE2 = re.compile(r'^(REsp\s*n\.|RE\s*n\.|AREsp\s*n\.)', re.IGNORECASE)
_RE_CABECALHO = re.compile(
    r'^([IVX]+\s*[–—\-]\s*[A-ZÁÉÍÓÚÂÊÔÃÕÇ]|[IVX]+\.\d+\s*[–—\-])'
)
_RE_ENDERECAMENTO = re.compile(r'^EXCELENTÍSSIM[OA]', re.IGNORECASE)
_RE_RAZOES_RECURSAIS = re.compile(r'^RAZÕES\s+RECURSAIS', re.IGNORECASE)
_RE_SAUDACAO = re.compile(
    r'^(EGRÉGI[OA]\s|COLENDAN?\s|EMINENTES?\s|EMÉRIT[OA]S?\s)',
    re.IGNORECASE,
)
_RE_QUALIFICACAO = re.compile(
    r'(por seu advogado|já qualificado nos autos|vem respeitosamente|'
    r'vem, respeitosamente|por intermedio de|por intermédio de)',
    re.IGNORECASE,
)
_RE_ALINEA = re.compile(r'^[a-z]\)\s')


# ── Utilitarios de template ───────────────────────────────────────────────────

def _dotx_para_docx_temp(dotx_path):
    """Copia .dotx para temp .docx com Content-Type corrigido para python-docx aceitar."""
    TIPO_TEMPLATE = (
        'application/vnd.openxmlformats-officedocument'
        '.wordprocessingml.template.main+xml'
    )
    TIPO_DOCUMENTO = (
        'application/vnd.openxmlformats-officedocument'
        '.wordprocessingml.document.main+xml'
    )
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.close()
    shutil.copy2(dotx_path, tmp.name)
    with zipfile.ZipFile(tmp.name, 'r') as zin:
        conteudos = {n: zin.read(n) for n in zin.namelist()}
    ct = conteudos['[Content_Types].xml'].decode('utf-8').replace(
        TIPO_TEMPLATE, TIPO_DOCUMENTO
    )
    conteudos['[Content_Types].xml'] = ct.encode('utf-8')
    with zipfile.ZipFile(tmp.name, 'w', zipfile.ZIP_DEFLATED) as zout:
        for nome, dados in conteudos.items():
            zout.writestr(nome, dados)
    return tmp.name


def _set_numId(para, num_id):
    """
    Define ou remove numeracao automatica de um paragrafo via OOXML.
    num_id=0  → sem numeracao
    num_id=1  → lista decimal 1. 2. 3. (paragrafos do corpo)
    num_id=2  → lista letras a) b) c) (alineas de pedidos)
    """
    pPr = para._p.get_or_add_pPr()
    numPr = pPr.find(qn('w:numPr'))
    if num_id == 0:
        if numPr is not None:
            pPr.remove(numPr)
        return
    if numPr is None:
        numPr = OxmlElement('w:numPr')
        pPr.append(numPr)
    ilvl = numPr.find(qn('w:ilvl'))
    if ilvl is None:
        ilvl = OxmlElement('w:ilvl')
        numPr.append(ilvl)
    ilvl.set(qn('w:val'), '0')
    nId = numPr.find(qn('w:numId'))
    if nId is None:
        nId = OxmlElement('w:numId')
        numPr.append(nId)
    nId.set(qn('w:val'), str(num_id))


def configurar_documento():
    """Carrega o template oficial e limpa o corpo para nova peca."""
    template_path = _TEMPLATE_PATH
    if os.path.exists(template_path):
        tmp_path = _dotx_para_docx_temp(template_path)
        doc = Document(tmp_path)
        os.unlink(tmp_path)
        # Limpar corpo mantendo sectPr (referencia cabecalho/rodape com timbre)
        body = doc.element.body
        for child in list(body):
            if not child.tag.endswith('}sectPr'):
                body.remove(child)
        print(f'Usando template: {template_path}')
    else:
        doc = Document()
        print('AVISO: template nao encontrado, gerando sem timbre.')
        for section in doc.sections:
            section.left_margin = Cm(3)
            section.right_margin = Cm(2)
            section.top_margin = Cm(3)
            section.bottom_margin = Cm(2)
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
    return doc


# ── Formatacao de texto ───────────────────────────────────────────────────────

def _run(para, texto, bold=None, italic=None, size_pt=None):
    """Adiciona run com formatacao ao paragrafo."""
    run = para.add_run(texto)
    run.font.name = 'Arial'
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    return run


def _markdown(para, texto, bold_base=False, italic_base=False, size_pt=12):
    """Processa **negrito** e *italico* inline e adiciona runs ao paragrafo."""
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)')
    last = 0
    for m in pattern.finditer(texto):
        if m.start() > last:
            _run(para, texto[last:m.start()], bold=bold_base, italic=italic_base, size_pt=size_pt)
        if m.group(2):
            _run(para, m.group(2), bold=True, italic=True, size_pt=size_pt)
        elif m.group(3):
            _run(para, m.group(3), bold=True, italic=italic_base, size_pt=size_pt)
        elif m.group(4):
            _run(para, m.group(4), bold=bold_base, italic=True, size_pt=size_pt)
        last = m.end()
    if last < len(texto):
        _run(para, texto[last:], bold=bold_base, italic=italic_base, size_pt=size_pt)


def _apenas_maiusculas(s):
    """Retorna True se a string e composta apenas por letras maiusculas (e pontuacao)."""
    letras = [c for c in s if c.isalpha()]
    return len(letras) > 3 and all(c.isupper() for c in letras)


# ── Processamento principal ───────────────────────────────────────────────────

def processar_texto(doc, texto):
    """
    Processa o texto da peca e aplica formatacao por tipo de paragrafo.

    Ordem de prioridade (maior para menor):
    1.  Linha vazia
    2.  Fechamento (Pede deferimento, data, assinatura)
    3.  Enderecamento (EXCELENTISSIMO)
    4.  Cabecalho de secao romano (I –, II –)
    5.  RAZOES RECURSAIS
    6.  Saudacao ao tribunal (EGREGIO, COLENDA, EMINENTES)
    7.  Linhas nao numeradas (Autos n., REsp n., Recorrente:, etc.)
    8.  Qualificacao das partes
    9.  Alinea a) b) c)
    10. Apenas maiusculas (nao capturado antes)
    11. Paragrafo normal (numerado via No Spacing)
    """
    for linha in texto.split('\n'):
        stripped = linha.strip()

        # 1. Linha vazia
        if not stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            continue

        # 2. Fechamento
        if _RE_FECHAMENTO.match(stripped):
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            _set_numId(p, 0)
            _markdown(p, stripped, size_pt=12)
            continue

        # 3. Enderecamento — grande espaco depois (replica Ariel REsp_vf)
        if _RE_ENDERECAMENTO.match(stripped):
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(54)   # ~3 linhas em branco
            _set_numId(p, 0)
            _run(p, stripped, size_pt=12)
            continue

        # 4. Cabecalho de secao romano — estilo Title (azul, negrito, do template)
        if _RE_CABECALHO.match(stripped):
            p = doc.add_paragraph(style='Title')
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            _set_numId(p, 0)
            p.add_run(stripped).font.name = 'Arial'
            continue

        # 5. RAZOES RECURSAIS — 13pt, centralizado, italico (padrao Ariel REsp_vf)
        if _RE_RAZOES_RECURSAIS.match(stripped):
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            _set_numId(p, 0)
            _run(p, stripped, bold=False, italic=True, size_pt=13)
            continue

        # 6. Saudacao ao tribunal — 13pt, centralizado, negrito
        if _RE_SAUDACAO.match(stripped):
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            _set_numId(p, 0)
            _run(p, stripped, bold=True, size_pt=13)
            continue

        # 7. Linhas nao numeradas (metadados, cabecalhos de identificacao)
        if _RE_NAO_NUMERADO.match(stripped):
            # "REsp n." / "RE n." iniciam Parte 2 — quebra de pagina antes
            if _RE_INICIO_PARTE2.match(stripped):
                p = doc.add_paragraph(style='Normal')
                p.paragraph_format.page_break_before = True
                p.paragraph_format.space_after = Pt(0)
                _set_numId(p, 0)
                _markdown(p, stripped, size_pt=11)
            elif re.match(r'^Autos\s*n\.', stripped, re.IGNORECASE):
                # "Autos n." na folha de rosto — espaco maior antes da qualificacao
                p = doc.add_paragraph(style='Normal')
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(36)  # ~2 linhas em branco
                _set_numId(p, 0)
                _markdown(p, stripped, size_pt=11)
            else:
                p = doc.add_paragraph(style='Normal')
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(4)
                _set_numId(p, 0)
                _markdown(p, stripped, size_pt=11)
            continue

        # 8. Qualificacao das partes
        if _RE_QUALIFICACAO.search(stripped):
            p = doc.add_paragraph(style='Qualificação das partes')
            p.paragraph_format.space_after = Pt(12)
            _set_numId(p, 0)
            _markdown(p, stripped, size_pt=12)
            continue

        # 9. Alinea a) b) c) (pedidos)
        if _RE_ALINEA.match(stripped):
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(4)
            _set_numId(p, 0)
            _markdown(p, stripped, size_pt=12)
            continue

        # 10. Apenas maiusculas (titulos nao detectados antes)
        if _apenas_maiusculas(stripped):
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            _set_numId(p, 0)
            _run(p, stripped, bold=True, size_pt=12)
            continue

        # 11. Paragrafo normal (numerado automaticamente via No Spacing)
        p = doc.add_paragraph(style='No Spacing')
        _markdown(p, stripped, size_pt=12)


# ── Helpers de path e data ────────────────────────────────────────────────────

def _sanitizar_nome(nome):
    """Remove acentos e caracteres especiais para nome de pasta/arquivo."""
    nfkd = unicodedata.normalize('NFKD', nome)
    sem_acento = ''.join(c for c in nfkd if not unicodedata.combining(c))
    sanitizado = re.sub(r'[^\w\s-]', '', sem_acento).strip()
    return re.sub(r'\s+', '_', sanitizado).upper()


def _construir_output_path(cliente, titulo, base_dir=None):
    """Constroi <base_dir>/<CLIENTE>/output_claude/<TIPO>_<CLIENTE>_<DATA>.docx"""
    if base_dir is None:
        base_dir = os.getcwd()
    nome_cliente = _sanitizar_nome(cliente)
    nome_titulo = _sanitizar_nome(titulo)
    data_hoje = datetime.now().strftime('%Y-%m-%d')
    nome_arquivo = f'{nome_titulo}_{nome_cliente}_{data_hoje}.docx'
    pasta = os.path.join(base_dir, nome_cliente, 'output_claude')
    os.makedirs(pasta, exist_ok=True)
    return os.path.join(pasta, nome_arquivo)


def _data_por_extenso():
    """Data atual por extenso no formato juridico brasileiro."""
    meses = {
        1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril',
        5: 'maio', 6: 'junho', 7: 'julho', 8: 'agosto',
        9: 'setembro', 10: 'outubro', 11: 'novembro', 12: 'dezembro',
    }
    hoje = datetime.now()
    return f'{hoje.day} de {meses[hoje.month]} de {hoje.year}'


def _conteudo_ja_tem_data(conteudo):
    """Retorna True se o conteudo ja contem bloco de encerramento (data/deferimento)."""
    ultimas = '\n'.join(conteudo.strip().split('\n')[-6:]).lower()
    return bool(
        re.search(r'\d{1,2}\s+de\s+\w+\s+de\s+\d{4}', ultimas)
        or re.search(r'pede\s+deferimento', ultimas)
    )


# ── Geracao do documento ──────────────────────────────────────────────────────

def gerar_peticao(titulo, conteudo, cliente, cidade='Florianópolis',
                  advogado=None, oab=None, base_dir=None):
    """Gera o .docx da peca usando o template oficial do escritorio."""
    output_path = _construir_output_path(cliente, titulo, base_dir)
    doc = configurar_documento()
    doc.core_properties.subject = f'Cliente: {cliente}'

    # Conteudo comeca direto no EXCELENTISSIMO — sem titulo antes
    processar_texto(doc, conteudo)

    # Bloco de encerramento automatico (se nao presente no conteudo)
    if not _conteudo_ja_tem_data(conteudo):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

        p_data = doc.add_paragraph(style='Normal')
        p_data.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_data.paragraph_format.space_before = Pt(12)
        p_data.paragraph_format.space_after = Pt(18)
        _set_numId(p_data, 0)
        _run(p_data, f'{cidade}, {_data_por_extenso()}.', size_pt=12)

        if advogado:
            p_adv = doc.add_paragraph(style='Normal')
            p_adv.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_adv.paragraph_format.space_before = Pt(24)
            p_adv.paragraph_format.space_after = Pt(2)
            _set_numId(p_adv, 0)
            _run(p_adv, advogado, bold=True, size_pt=12)

            if oab:
                p_oab = doc.add_paragraph(style='Normal')
                p_oab.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_oab.paragraph_format.space_before = Pt(0)
                p_oab.paragraph_format.space_after = Pt(6)
                _set_numId(p_oab, 0)
                _run(p_oab, f'OAB/SC n. {oab}', size_pt=12)

    doc.save(output_path)
    print(f'Petição gerada: {output_path}')
    return output_path


# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    sys.stdout.reconfigure(encoding='utf-8')
    parser = argparse.ArgumentParser(
        description=(
            'Gera petição jurídica em .docx com timbre do escritório. '
            'Salva em <base-dir>/<CLIENTE>/output_claude/<TIPO>_<CLIENTE>_<DATA>.docx'
        )
    )
    parser.add_argument('--titulo', required=True, help='Tipo da peca (ex: CONTESTACAO)')
    parser.add_argument('--cliente', required=True, help='Nome do cliente')
    parser.add_argument('--conteudo', help='Caminho para arquivo .txt com o conteudo')
    parser.add_argument('--base-dir', default=None, help='Diretorio base para salvar')
    parser.add_argument('--cidade', default='Florianópolis')
    parser.add_argument('--advogado', default=None)
    parser.add_argument('--oab', default=None)
    args = parser.parse_args()

    if args.conteudo:
        with open(args.conteudo, 'r', encoding='utf-8') as f:
            conteudo = f.read()
    else:
        conteudo = sys.stdin.read()

    if not conteudo.strip():
        print('Erro: nenhum conteudo fornecido.', file=sys.stderr)
        sys.exit(1)

    gerar_peticao(
        titulo=args.titulo,
        conteudo=conteudo,
        cliente=args.cliente,
        cidade=args.cidade,
        advogado=args.advogado,
        oab=args.oab,
        base_dir=args.base_dir,
    )


if __name__ == '__main__':
    main()
